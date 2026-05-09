"""Resume tools for parsing and writing resume files."""

import os
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_resume_text(file_path: str) -> str:
    """Extract text from a resume file (PDF or DOCX).
    
    Args:
        file_path: Path to the resume file (.pdf or .docx)
        
    Returns:
        Extracted text content from the resume
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension == ".pdf":
        return _extract_from_pdf(path)
    elif extension == ".docx":
        return _extract_from_docx(path)
    else:
        raise ValueError(f"Unsupported resume format: {extension}. Only .pdf and .docx are supported.")


def _extract_from_pdf(path: Path) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_parts = []
    
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    return "\n\n".join(text_parts)


def _extract_from_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    return "\n\n".join(paragraphs)


def create_resume_docx(
    content: str,
    output_path: str,
    font_name: str = "Arial",
    font_size: int = 11,
) -> None:
    """Create a DOCX resume file from text content.
    
    Args:
        content: Resume content to write (supports simple markdown-like formatting)
        output_path: Path where the DOCX file should be saved
        font_name: Font to use for the document
        font_size: Base font size in points
    """
    doc = Document()
    
    # Set default style
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    
    # Parse content into sections
    sections = _parse_resume_content(content)
    
    for section in sections:
        if section["type"] == "heading":
            # Add heading
            heading = doc.add_heading(section["content"], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.name = font_name
                run.font.size = Pt(14)
                run.font.bold = True
        elif section["type"] == "subheading":
            # Add subheading (e.g., company name, date range)
            para = doc.add_paragraph()
            run = para.add_run(section["content"])
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = True
        elif section["type"] == "bullet":
            # Add bullet point
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(section["content"])
            run.font.name = font_name
            run.font.size = Pt(font_size)
        else:
            # Regular paragraph
            para = doc.add_paragraph(section["content"])
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Ensure output directory exists
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(output_path)


def _parse_resume_content(content: str) -> list[dict]:
    """Parse resume content into structured sections.
    
    Supports:
    - Headings: Lines starting with # or all caps
    - Subheadings: Lines starting with -
    - Bullets: Lines starting with •
    - Regular paragraphs
    
    Args:
        content: Raw resume content
        
    Returns:
        List of section dictionaries with type and content
    """
    sections = []
    lines = content.split("\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for heading markers
        if line.startswith("# "):
            sections.append({"type": "heading", "content": line[2:]})
        elif line.startswith("## "):
            sections.append({"type": "heading", "content": line[3:]})
        elif line.startswith("- "):
            sections.append({"type": "subheading", "content": line[2:]})
        elif line.startswith("• "):
            sections.append({"type": "bullet", "content": line[2:]})
        elif _is_all_caps(line):
            sections.append({"type": "heading", "content": line})
        else:
            sections.append({"type": "paragraph", "content": line})
    
    return sections


def _is_all_caps(text: str) -> bool:
    """Check if text is in ALL CAPS (considering punctuation and spaces)."""
    import re
    # Remove non-letter characters for check
    letters = re.sub(r"[^a-zA-Z]", "", text)
    return len(letters) > 0 and letters.isupper()


def get_file_format(file_path: str) -> str:
    """Get the file format from extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        "pdf", "docx", or "unknown"
    """
    extension = Path(file_path).suffix.lower()
    
    if extension == ".pdf":
        return "pdf"
    elif extension == ".docx":
        return "docx"
    else:
        return "unknown"


def validate_resume_file(file_path: str) -> tuple[bool, str]:
    """Validate that a resume file exists and has a supported format.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    path = Path(file_path)
    
    if not path.exists():
        return False, f"File not found: {file_path}"
    
    extension = path.suffix.lower()
    
    if extension not in [".pdf", ".docx"]:
        return False, f"Unsupported format: {extension}. Only .pdf and .docx are supported."
    
    return True, ""