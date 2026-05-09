"""Job Description tools for managing local job description files."""

import os
import json
from pathlib import Path
from typing import Optional, List
from datetime import datetime


# Default path for job descriptions
JOB_DESCRIPTIONS_PATH = Path("data/job-descriptions")


def get_job_descriptions_path() -> Path:
    """Get the path to the job descriptions directory."""
    return JOB_DESCRIPTIONS_PATH


def list_job_descriptions() -> List[dict]:
    """List all job descriptions in the job-descriptions folder.
    
    Returns:
        List of dictionaries with job description metadata
    """
    path = get_job_descriptions_path()
    
    if not path.exists():
        return []
    
    job_descriptions = []
    
    for file_path in path.iterdir():
        if file_path.is_file() and file_path.suffix in [".txt", ".md", ".json"]:
            stat = file_path.stat()
            
            # Try to extract job info from filename
            filename = file_path.stem
            
            job_desc = {
                "id": filename,
                "filename": file_path.name,
                "path": str(file_path),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            }
            
            # Try to get title from first line if markdown/txt
            if file_path.suffix in [".txt", ".md"]:
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        first_line = f.readline().strip()
                        if first_line:
                            job_desc["title"] = first_line
                except Exception:
                    pass
            
            # For JSON, try to parse metadata
            elif file_path.suffix == ".json":
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        job_desc["title"] = data.get("title", filename)
                        job_desc["company"] = data.get("company", "")
                        job_desc["location"] = data.get("location", "")
                except Exception:
                    pass
            
            job_descriptions.append(job_desc)
    
    return job_descriptions


def read_job_description(job_id: str) -> Optional[dict]:
    """Read a job description by its ID (filename without extension).
    
    Args:
        job_id: The job description ID (filename without extension)
        
    Returns:
        Dictionary with job description data, or None if not found
    """
    path = get_job_descriptions_path()
    
    # Try different extensions
    for ext in [".txt", ".md", ".json"]:
        file_path = path / f"{job_id}{ext}"
        if file_path.exists():
            return _read_job_file(file_path)
    
    return None


def _read_job_file(file_path: Path) -> dict:
    """Read a job description file and return its contents."""
    result = {
        "id": file_path.stem,
        "filename": file_path.name,
        "path": str(file_path),
        "format": file_path.suffix.lstrip("."),
    }
    
    if file_path.suffix == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            result["title"] = data.get("title", "")
            result["company"] = data.get("company", "")
            result["location"] = data.get("location", "")
            result["description"] = data.get("description", "")
            result["requirements"] = data.get("requirements", [])
            result["raw_data"] = data
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            lines = content.split("\n")
            
            # First line is often the title
            result["title"] = lines[0].strip() if lines else ""
            result["description"] = content
            result["raw_data"] = content
    
    return result


def save_job_description(
    job_id: str,
    content: str,
    format: str = "txt",
    metadata: Optional[dict] = None
) -> str:
    """Save a job description to a file.
    
    Args:
        job_id: Unique identifier for the job (used as filename)
        content: Job description content
        format: File format ("txt", "md", or "json")
        metadata: Optional metadata dict (for JSON format)
        
    Returns:
        Path to the saved file
    """
    path = get_job_descriptions_path()
    path.mkdir(parents=True, exist_ok=True)
    
    file_path = path / f"{job_id}.{format}"
    
    if format == "json":
        data = metadata or {}
        data["description"] = content
        data["title"] = metadata.get("title", job_id) if metadata else job_id
        
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    else:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
    
    return str(file_path)


def search_job_descriptions(query: str) -> List[dict]:
    """Search job descriptions by content or title.
    
    Args:
        query: Search query string
        
    Returns:
        List of matching job description dicts with relevance info
    """
    job_list = list_job_descriptions()
    query_lower = query.lower()
    
    results = []
    
    for job in job_list:
        # Read the full content
        job_data = read_job_description(job["id"])
        if not job_data:
            continue
        
        content = job_data.get("description", "").lower()
        title = job_data.get("title", "").lower()
        
        # Simple relevance scoring
        score = 0
        if query_lower in title:
            score += 10
        if query_lower in content:
            score += 5
        
        # Check individual words
        query_words = query_lower.split()
        for word in query_words:
            if len(word) > 2:
                if word in title:
                    score += 3
                if word in content:
                    score += 1
        
        if score > 0:
            job["relevance_score"] = score
            job["matched_content"] = _get_matched_snippet(content, query_lower)
            results.append(job)
    
    # Sort by relevance
    results.sort(key=lambda x: x.get("relevance_score", 0), reverse=True)
    
    return results


def _get_matched_snippet(content: str, query: str, context_chars: int = 100) -> str:
    """Extract a snippet around the matched query."""
    query_lower = query.lower()
    content_lower = content.lower()
    
    pos = content_lower.find(query_lower)
    if pos == -1:
        # Try to find any word from query
        words = query.split()
        for word in words:
            if len(word) > 3:
                pos = content_lower.find(word.lower())
                if pos != -1:
                    break
        
        if pos == -1:
            return content[:context_chars * 2] + "..."
    
    # Get context around match
    start = max(0, pos - context_chars)
    end = min(len(content), pos + len(query) + context_chars)
    
    snippet = content[start:end]
    
    if start > 0:
        snippet = "..." + snippet
    if end < len(content):
        snippet = snippet + "..."
    
    return snippet


def delete_job_description(job_id: str) -> bool:
    """Delete a job description file.
    
    Args:
        job_id: The job description ID to delete
        
    Returns:
        True if deleted, False if not found
    """
    path = get_job_descriptions_path()
    
    for ext in [".txt", ".md", ".json"]:
        file_path = path / f"{job_id}{ext}"
        if file_path.exists():
            file_path.unlink()
            return True
    
    return False


def get_job_description_by_keyword(keyword: str) -> Optional[dict]:
    """Find a job description by keyword in title or filename.
    
    Args:
        keyword: Keyword to search for
        
    Returns:
        First matching job description or None
    """
    # First search by filename
    job_id = keyword.replace(" ", "_").lower()
    job = read_job_description(job_id)
    if job:
        return job
    
    # Search by content
    results = search_job_descriptions(keyword)
    if results:
        job_id = results[0]["id"]
        return read_job_description(job_id)
    
    return None


def create_job_from_text(
    title: str,
    company: str,
    location: str,
    description: str,
    requirements: Optional[List[str]] = None
) -> str:
    """Create a new job description from text components.
    
    Args:
        title: Job title
        company: Company name
        location: Job location
        description: Job description text
        requirements: Optional list of requirements
        
    Returns:
        Path to the created file
    """
    import uuid
    
    # Create a slug for the job ID
    job_id = f"{title.lower().replace(' ', '_')[:30]}_{uuid.uuid4().hex[:8]}"
    
    metadata = {
        "title": title,
        "company": company,
        "location": location,
        "requirements": requirements or [],
    }
    
    return save_job_description(job_id, description, format="json", metadata=metadata)


def extract_text_from_file(file_path: str) -> str:
    """Extract text from a file (PDF, DOCX, TXT, or MD).
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension == ".pdf":
        return _extract_text_from_pdf(path)
    elif extension == ".docx":
        return _extract_text_from_docx(path)
    elif extension in [".txt", ".md"]:
        return _extract_text_from_txt(path)
    else:
        raise ValueError(f"Unsupported file format: {extension}. Supported formats: .pdf, .docx, .txt, .md")


def _extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber
    text_parts = []
    
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    return "\n\n".join(text_parts)


def _extract_text_from_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    paragraphs = []
    
    doc = Document(path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    return "\n\n".join(paragraphs)


def _extract_text_from_txt(path: Path) -> str:
    """Extract text from a TXT or MD file."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def upload_job_description_file(
    file_path: str,
    title: Optional[str] = None,
    company: Optional[str] = None,
    location: Optional[str] = None
) -> dict:
    """Upload and parse a job description from a file (PDF, DOCX, TXT, or MD).
    
    This function extracts text from the file, generates a job ID from the filename,
    and saves it to the job-descriptions directory.
    
    Args:
        file_path: Path to the job description file (.pdf, .docx, .txt, .md)
        title: Optional job title (if not provided, extracted from filename or first line)
        company: Optional company name
        location: Optional job location
        
    Returns:
        Dictionary with job description data including:
        - id: Unique job identifier
        - title: Job title
        - company: Company name
        - location: Job location
        - description: Full job description text
        - path: Path to the saved file
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    import uuid
    
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension not in [".pdf", ".docx", ".txt", ".md"]:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: .pdf, .docx, .txt, .md"
        )
    
    # Extract text based on file type
    description = extract_text_from_file(file_path)
    
    # Generate job_id from filename
    filename = path.stem
    job_id = f"{filename.lower().replace(' ', '_')[:30]}_{uuid.uuid4().hex[:8]}"
    
    # Use provided title or extract from content
    if not title:
        # Try to get title from first non-empty line
        lines = description.strip().split("\n")
        for line in lines:
            line = line.strip()
            if line and len(line) > 3:
                title = line[:100]  # Limit title length
                break
        if not title:
            title = filename
    
    # Save the job description
    metadata = {
        "title": title,
        "company": company or "",
        "location": location or "",
    }
    
    saved_path = save_job_description(job_id, description, format="json", metadata=metadata)
    
    return {
        "id": job_id,
        "title": title,
        "company": company or "",
        "location": location or "",
        "description": description,
        "path": saved_path,
    }


def validate_job_description_file(file_path: str) -> tuple[bool, str]:
    """Validate that a job description file exists and has a supported format.
    
    Args:
        file_path: Path to the job description file
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    path = Path(file_path)
    
    if not path.exists():
        return False, f"File not found: {file_path}"
    
    extension = path.suffix.lower()
    
    if extension not in [".pdf", ".docx", ".txt", ".md"]:
        return False, f"Unsupported format: {extension}. Supported: .pdf, .docx, .txt, .md"
    
    return True, ""