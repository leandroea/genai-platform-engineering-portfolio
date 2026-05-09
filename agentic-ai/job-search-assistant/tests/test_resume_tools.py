"""Tests for resume tools."""

import pytest
from pathlib import Path
import tempfile
import os

from src.tools.resume_tools import (
    extract_resume_text,
    create_resume_docx,
    get_file_format,
    validate_resume_file,
    _is_all_caps,
    _parse_resume_content,
)


class TestResumeTools:
    """Test suite for resume_tools module."""
    
    def test_get_file_format(self):
        """Test file format detection."""
        assert get_file_format("resume.pdf") == "pdf"
        assert get_file_format("resume.docx") == "docx"
        assert get_file_format("resume.txt") == "unknown"
        assert get_file_format("/path/to/resume.PDF") == "pdf"
    
    def test_is_all_caps(self):
        """Test ALL CAPS detection."""
        assert _is_all_caps("EXPERIENCE") == True
        assert _is_all_caps("SKILLS") == True
        assert _is_all_caps("John Smith") == False
        assert _is_all_caps("Python Developer") == False
        assert _is_all_caps("") == False
    
    def test_parse_resume_content(self):
        """Test resume content parsing."""
        content = """# John Smith
Software Engineer

## Experience
- Senior Developer at TechCorp
• Led team of 5 engineers

Skills
Python, JavaScript, AWS"""
        
        sections = _parse_resume_content(content)
        
        assert len(sections) > 0
        # Check that headings are detected
        headings = [s for s in sections if s["type"] == "heading"]
        assert len(headings) >= 1
    
    def test_validate_resume_file_nonexistent(self):
        """Test validation of non-existent file."""
        is_valid, error = validate_resume_file("nonexistent.pdf")
        assert is_valid == False
        assert "not found" in error.lower()
    
    def test_validate_resume_file_unsupported_format(self):
        """Test validation of unsupported format."""
        # Create a temporary file with wrong extension
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"test content")
            temp_path = f.name
        
        try:
            is_valid, error = validate_resume_file(temp_path)
            assert is_valid == False
            assert "unsupported" in error.lower()
        finally:
            os.unlink(temp_path)
    
    def test_create_resume_docx(self):
        """Test DOCX resume creation."""
        content = """# John Smith
Software Engineer

## Experience
- Senior Developer at TechCorp
• Led team of 5 engineers
• Increased revenue by 30%

## Skills
Python, JavaScript, AWS"""
        
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "test_resume.docx")
            
            create_resume_docx(content, output_path)
            
            assert Path(output_path).exists()
            assert Path(output_path).stat().st_size > 0
    
    def test_create_resume_docx_with_custom_formatting(self):
        """Test DOCX creation with custom formatting."""
        content = "# Header\n\nSome text\n\n## Subheader\n- Bullet 1\n- Bullet 2"
        
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "styled_resume.docx")
            
            create_resume_docx(content, output_path, font_name="Times New Roman", font_size=12)
            
            assert Path(output_path).exists()


class TestResumeParsing:
    """Test suite for resume text extraction (requires real files)."""
    
    @pytest.fixture
    def sample_pdf_path(self):
        """Create a sample PDF for testing."""
        # Note: In a real scenario, you'd have an actual PDF file
        # For unit testing without real files, we skip actual parsing
        pytest.skip("Requires real PDF file")
    
    @pytest.fixture
    def sample_docx_path(self):
        """Create a sample DOCX for testing."""
        # Create a temporary DOCX file
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "sample.docx")
            
            doc = Document()
            doc.add_heading("John Smith", 0)
            doc.add_paragraph("Software Engineer")
            doc.add_heading("Experience", level=1)
            doc.add_paragraph("Worked at TechCorp")
            doc.save(docx_path)
            
            yield docx_path
    
    def test_extract_from_docx(self, sample_docx_path):
        """Test text extraction from DOCX."""
        text = extract_resume_text(sample_docx_path)
        
        assert "John Smith" in text
        assert "Software Engineer" in text
        assert "TechCorp" in text
    
    def test_extract_from_pdf_requires_real_file(self):
        """Test that PDF extraction requires a real file."""
        with pytest.raises(FileNotFoundError):
            extract_resume_text("nonexistent.pdf")
    
    def test_extract_unsupported_format(self):
        """Test extraction from unsupported format raises error."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"text content")
            temp_path = f.name
        
        try:
            with pytest.raises(ValueError) as exc_info:
                extract_resume_text(temp_path)
            assert "unsupported" in str(exc_info.value).lower()
        finally:
            os.unlink(temp_path)